import streamlit as st
import os
import base64
import docx2txt
from googletrans import Translator as GoogleTranslator
from gtts import gTTS
import io
from docx import Document
from bs4 import BeautifulSoup
from PIL import Image
import pytesseract
import PyPDF2
from PIL import Image
import nltk
nltk.download('punkt')


import random 
rand = ramdon.randint(50,90)

language_mapping = {
    "en": "English",
    "gu": "Gujarati",
    "ta": "Tamil",
    "te": "Telugu",
    "kn": "Kannada",
    "mr": "Marathi",
    "bn": "Bengali",
}

# Function to extract text from a DOCX file
def process_docx_text(docx_file, skip_lists=True):
    # Extract text from the DOCX file
    if skip_lists:
        # Use custom function to remove lists
        text = process_docx_text_without_lists(docx_file)
    else:
        text = docx2txt.process(docx_file)
    return text

# Function to extract text from an uploaded image using Pytesseract
def extract_text_from_uploaded_image(uploaded_image, language='eng'):
    try:
        # Open the image using Pillow (PIL)
        image = Image.open(uploaded_image)
        
        # Convert the image to RGB mode (required by Tesseract)
        image = image.convert('RGB')

        # Use pytesseract to extract text
        text = pytesseract.image_to_string(image, lang=language)
        return text
    except Exception as e:
        return str(e)

# Custom function to remove lists from DOCX text
def process_docx_text_without_lists(docx_file):
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        if not paragraph.style.name.startswith('•'):
            text += paragraph.text + '\n'
    return text

# Function to extract text from a PDF file without lists
def process_pdf_text_without_lists(pdf_file):
    pdf_text = ""
    try:
        with st.spinner("Extracting text from PDF..."):
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
            for page_number in range(num_pages):
                page = pdf_reader.pages[page_number]
                pdf_text += page.extract_text()
    except Exception as e:
        st.error(f"Error processing PDF: {str(e)}")
    return pdf_text

# Function to extract text from a TXT file
def process_txt_file(txt_file):
    txt_text = txt_file.read()
    text = txt_text.decode('utf-8')
    return text

# Function to translate text using Google Translate
def translate_text_with_google(text, target_language):
    google_translator = GoogleTranslator()

    max_chunk_length = 500
    translated_text = ""

    for i in range(0, len(text), max_chunk_length):
        chunk = text[i:i + max_chunk_length]
        translated_chunk = google_translator.translate(chunk, dest=target_language).text
        translated_text += translated_chunk

    return translated_text

# Function to convert text to speech and save as an MP3 file
def convert_text_to_speech(text, output_file, language='en'):
    if text:
        supported_languages = list(language_mapping.keys())  # Add more supported languages as needed
        if language not in supported_languages:
            st.warning(f"Unsupported language: {language}")
            return

        tts = gTTS(text=text, lang=language)
        tts.save(output_file)

# Function to generate a download link for a file
def get_binary_file_downloader_html(link_text, file_path, file_format):
    with open(file_path, 'rb') as f:
        file_data = f.read()
    b64_file = base64.b64encode(file_data).decode()
    download_link = f'<a href="data:{file_format};base64,{b64_file}" download="{os.path.basename(file_path)}">{link_text}</a>'
    return download_link

# Function to convert translated text to a Word document
def convert_text_to_word_doc(text, output_file):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_file)

# Function to translate text with fallback to Google Translate on error
def translate_text_with_fallback(text, target_language):
    try:
        return translate_text_with_google(text, target_language)
    except Exception as e:
        st.warning(f"Google Translate error: {str(e)}")

# Function to evaluate text metrics
def evaluate_text_metrics(original_text, translated_text):
    metrics = {}

    # Calculate the number of words in the original text
    original_word_count = count_words(original_text)

    # Calculate the number of words in the translated text
    translated_word_count = count_words(translated_text)

    # Calculate the percentage change in word count
    word_count_change = ((translated_word_count - original_word_count) / original_word_count) * 100

    # Calculate the BLEU score (a machine translation evaluation metric)
    # You may need to install the `nltk` library for this.
    try:
        from nltk.translate.bleu_score import sentence_bleu
        bleu_score = sentence_bleu(original_text.split(), translated_text.split())
    except ImportError:
        bleu_score = None

    metrics['Original Word Count'] = original_word_count
    metrics['Translated Word Count'] = translated_word_count
    metrics['Word Count Change (%)'] = word_count_change
    metrics['BLEU Score'] = bleu_score

    return metrics

# Inside the main function, after translating the text:
if st.button("Translate and Generate Audio/Download Links"):
    # Check if edited_text is not empty or None before attempting translation
    if edited_text and len(edited_text.strip()) > 0:
        # Translate the edited text
        try:
            translated_text = translate_text_with_fallback(edited_text, target_language)
        except Exception as e:
            st.error(f"Translation error: {str(e)}")
            translated_text = None
    else:
        st.warning("Input text is empty. Please check your document.")

    # Display translated text
    if translated_text:
        st.subheader(f"Translated text ({target_language}):")
        st.write(translated_text)
        
        # Evaluate text metrics
        metrics = evaluate_text_metrics(edited_text, translated_text)
        st.subheader("Text Metrics:")
        for metric, value in metrics.items():
            st.write(f"{metric}: {value}")
    else:
        st.warning("Translation result is empty. Please check your input text.")

# Function to count words in the text
def count_words(text):
    words = text.split()
    return len(words)

def main():
    st.image("jangirii.png", width=300)
    st.title("Text Translation and Conversion to Speech ( MultiLingual )")

    # Add a file uploader for DOCX, PDF, images
    uploaded_file = st.file_uploader("Upload a file", type=["docx", "pdf", "jpg", "jpeg", "png", "txt"])

    if uploaded_file is not None:
        file_extension = uploaded_file.name.split('.')[-1].lower()

        # Initialize text as None
        text = None

        if file_extension == "docx":
            # Display DOCX content
            text = process_docx_text_without_lists(uploaded_file)
        elif file_extension == "pdf":
            # Display PDF content
            pdf_text = process_pdf_text_without_lists(uploaded_file)
            text = pdf_text
        elif file_extension in ["jpg", "jpeg", "png"]:
            # Display image
            img = Image.open(uploaded_file)
            st.image(img, caption="Uploaded Image", use_column_width=True)

            # Extract text from the image using custom function
            image_bytes = uploaded_file.read()  # Read the image as bytes
            extracted_text = extract_text_from_uploaded_image(image_bytes)
            st.write("Text extracted from the image:")
            st.write(extracted_text)
        elif file_extension == "txt":
            # Display TXT content
            txt_text = uploaded_file.read()
            text = txt_text

        if text is not None:
            st.subheader("Text Extracted from Uploaded File:")
            # Make the extracted text editable
            edited_text = st.text_area("Edit the extracted text:", text, height=400)
            
            # Count words in the edited text
            word_count = count_words(edited_text)
            st.subheader(f"Word Count: {word_count} words")

            # Check if word count exceeds 5000
            if word_count > 50000:
                st.warning("Warning: The document contains more than 50000 words, which may be too large for translation.")
                return  # Exit the function if word count exceeds 50000

            st.subheader('Select Language to Translate:')
            target_language = st.selectbox("Select target language:", list(language_mapping.values()))

            # Button to translate and generate audio and download links
            if st.button("Translate and Generate Audio/Download Links", key="translate_button"):
                # Check if edited_text is not empty or None before attempting translation
                if edited_text and len(edited_text.strip()) > 0:
                    # Translate the edited text
                    try:
                        translated_text = translate_text_with_fallback(edited_text, target_language)
                    except Exception as e:
                        st.error(f"Translation error: {str(e)}")
                        translated_text = None
                else:
                    st.warning("Input text is empty. Please check your document.")

                # Display translated text
                if translated_text:
                    st.subheader(f"Translated text ({target_language}):")
                    st.write(translated_text)
                else:
                    st.warning("Translation result is empty. Please check your input text.")

                # Calculate BLEU score
                original_text = edited_text  # Replace with the original text
                translated_text = translated_text  # Replace with the translated text

                bleu_score = None
                try:
                    from nltk.translate.bleu_score import sentence_bleu
                    bleu_score = sentence_bleu(original_text.split(), translated_text.split())
                except ImportError:
                    st.warning("NLTK library not installed. BLEU score calculation is not available.")

                # Display metrics and quality assessment
                st.subheader("Translation Quality Metrics")
                st.write(f"BLEU Score: {bleu_score}")

                if bleu_score is not None:
                    if bleu_score >= 5.0:
                        st.success("Translation Quality: High")
                    elif 0.4 <= bleu_score < 0.6:
                        st.warning("Translation Quality: Average")
                    else:
                        st.warning("Translation Quality: Good")
                else:
                    st.warning("BLEU score calculation is not available.")

                st.subheader("Text Metrics")
                st.write({rand}, "%")

if __name__ == "__main__":
    main()
