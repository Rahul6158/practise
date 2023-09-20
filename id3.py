import streamlit as st
import os
import base64
import docx2txt
from translate import Translator
from gtts import gTTS
import io
from docx import Document
from bs4 import BeautifulSoup
from PIL import Image
import PyPDF2
import pytesseract
import easyocr
from PIL import Image
import speech_recognition as sr
from pydub import AudioSegment
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
import heapq

# Function to extract text from a DOCX file
def process_docx_text(docx_file, skip_lists=True):
    # Extract text from the DOCX file
    if skip_lists:
        # Use custom function to remove lists
        text = process_docx_text_without_lists(docx_file)
    else:
        text = docx2txt.process(docx_file)
    return text

# Function to extract text from an image using easyocr
def extract_text_from_image(image_bytes):
    try:
        image = Image.open(io.BytesIO(image_bytes))
        reader = easyocr.Reader(['en'])  # Specify the language(s) you want to recognize
        results = reader.readtext(image)

        text = ""
        for (bbox, text, prob) in results:
            text += text + " "

        return text.strip()
    except Exception as e:
        st.error(f"Error extracting text from image: {str(e)}")
        return ""


# Custom function to remove lists from DOCX text
def process_docx_text_without_lists(docx_file):
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        if not paragraph.style.name.startswith('•'):
            text += paragraph.text + '\n'
    return text

# Function to extract text from a PDF file without lists
def process_pdf_text_without_lists(pdf_file):
    pdf_text = ""
    try:
        with st.spinner("Extracting text from PDF..."):
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
            for page_number in range(num_pages):
                page = pdf_reader.pages[page_number]
                pdf_text += page.extract_text()
    except Exception as e:
        st.error(f"Error processing PDF: {str(e)}")
    return pdf_text

# Function to translate text using the translate library with a loop
def translate_text(text, target_language):
    translator = Translator(to_lang=target_language)
    max_chunk_length = 500
    translated_text = ""

    for i in range(0, len(text), max_chunk_length):
        chunk = text[i:i + max_chunk_length]
        translated_chunk = translator.translate(chunk)
        translated_text += translated_chunk

    return translated_text


# Function to convert text to speech and save as an MP3 file
def convert_text_to_speech(text, output_file, language='en'):
    if text:
        tts = gTTS(text=text, lang=language)
        tts.save(output_file)

# Function to generate a download link for a file
def get_binary_file_downloader_html(link_text, file_path, file_format):
    with open(file_path, 'rb') as f:
        file_data = f.read()
    b64_file = base64.b64encode(file_data).decode()
    download_link = f'<a href="data:{file_format};base64,{b64_file}" download="{os.path.basename(file_path)}">{link_text}</a>'
    return download_link

# Function to convert translated text to a Word document
def convert_text_to_word_doc(text, output_file):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_file)

# Function to convert Word document to HTML
def convert_word_doc_to_html(docx_file):
    txt = docx2txt.process(docx_file)
    soup = BeautifulSoup(txt, 'html.parser')
    return soup.prettify()

# Function to translate text using the translate library with a loop
def translate_text(text, target_language):
    translator = Translator(to_lang=target_language)
    max_chunk_length = 500
    translated_text = ""

    for i in range(0, len(text), max_chunk_length):
        chunk = text[i:i + max_chunk_length]
        translated_chunk = translator.translate(chunk)
        translated_text += translated_chunk

    return translated_text

# Function to translate text using Google Translate with a loop
def translate_text_with_google(text, target_language):
    google_translator = GoogleTranslator()

    max_chunk_length = 500
    translated_text = ""

    for i in range(0, len(text), max_chunk_length):
        chunk = text[i:i + max_chunk_length]
        translated_chunk = google_translator.translate(chunk, dest=target_language).text
        translated_text += translated_chunk

    return translated_text

# Function to translate text with fallback to Google Translate on errors
def translate_text_with_fallback(text, target_language):
    try:
        return translate_text(text, target_language)
    except Exception as e:
        st.warning(f"MyMemory translation error: {str(e)}")

    # If MyMemory fails, use Google Translate
    st.warning("Falling back to Google Translate...")
    return translate_text_with_google(text, target_language)

# Function to recognize speech from an audio file and return text
def convert_audio_to_wav(audio_bytes):
    try:
        # Load audio data from bytes
        audio = AudioSegment.from_file(io.BytesIO(audio_bytes))
        
        # Export audio as WAV format
        wav_audio = audio.export(format="wav")
        
        # Convert to bytes
        wav_bytes = wav_audio.read()
        
        return wav_bytes
    except Exception as e:
        st.error(f"Error converting audio to WAV format: {str(e)}")
        return None

def summarize_large_text(text, num_sentences=5):
    # Tokenize the text into sentences
    sentences = sent_tokenize(text)

    # Tokenize the text into words
    words = word_tokenize(text)

    # Remove stopwords and punctuation
    stop_words = set(stopwords.words("english"))
    filtered_words = [word.lower() for word in words if word.isalnum() and word.lower() not in stop_words]

    # Calculate word frequency
    word_freq = FreqDist(filtered_words)

    # Calculate sentence scores based on word frequency
    sentence_scores = {}
    for sentence in sentences:
        for word in word_tokenize(sentence.lower()):
            if word in word_freq:
                if sentence not in sentence_scores:
                    sentence_scores[sentence] = word_freq[word]
                else:
                    sentence_scores[sentence] += word_freq[word]

    # Get the top N sentences with the highest scores
    summary_sentences = heapq.nlargest(num_sentences, sentence_scores, key=sentence_scores.get)

    # Combine the selected sentences into the summary
    summary = " ".join(summary_sentences)

    return summary


# Function to count words in the text
def count_words(text):
    words = text.split()
    return len(words)

language_mapping = {
    "en": "English",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
    "it": "Italian",
    "pt": "Portuguese",
    "nl": "Dutch",
    "hi": "Hindi",
    "ja": "Japanese",
    "ko": "Korean",
    "zh-cn": "Simplified Chinese",
    "ru": "Russian",
    "ar": "Arabic",
    "th": "Thai",
    "tr": "Turkish",
    "pl": "Polish",
    "cs": "Czech",
    "sv": "Swedish",
    "da": "Danish",
    "fi": "Finnish",
    "el": "Greek",
    "hu": "Hungarian",
    "uk": "Ukrainian",
    "no": "Norwegian",
    "id": "Indonesian",
    "vi": "Vietnamese",
    "ro": "Romanian",
    "bn": "Bengali",
    "fa": "Persian",
    "iw": "Hebrew",
    "bg": "Bulgarian",
    "ca": "Catalan",
    "hr": "Croatian",
    "sr": "Serbian",
    "sk": "Slovak",
    "sl": "Slovenian",
    "lt": "Lithuanian",
    "lv": "Latvian",
    "et": "Estonian",
    "is": "Icelandic",
    "ga": "Irish",
    "sq": "Albanian",
    "mk": "Macedonian",
    "hy": "Armenian",
    "ka": "Georgian",
    "mt": "Maltese",
    "mr": "Marathi",
    "ta": "Tamil",
    "te": "Telugu",
    "ur": "Urdu",
    "ne": "Nepali",
    "si": "Sinhala",
    "km": "Khmer",
    "lo": "Lao",
    "my": "Burmese",
    "jw": "Javanese",
    "mn": "Mongolian",
    "zu": "Zulu",
    "xh": "Xhosa"
}

# Main Streamlit app
# Main Streamlit app
def main():
    st.image("jangirii.png", width=50)
    st.title("Text Translation and Conversion to Speech (English - other languages)")

    # Add a file uploader for DOCX, PDF, images, and audio
    uploaded_file = st.file_uploader("Upload a file", type=["docx", "pdf", "jpg", "jpeg", "png", "txt", "mp3"])

    if uploaded_file is not None:
        file_extension = uploaded_file.name.split('.')[-1].lower()

        # Initialize text as None
        text = None

        if file_extension == "docx":
            # Display DOCX content
            text = process_docx_text_without_lists(uploaded_file)
        elif file_extension == "pdf":
            # Display PDF content
            pdf_text = process_pdf_text_without_lists(uploaded_file)
            text = pdf_text
        elif file_extension in ["jpg", "jpeg", "png"]:
            # Display image
            img = Image.open(uploaded_file)
            st.image(img, caption="Uploaded Image", use_column_width=True)

            # Extract text from the image using custom function
            image_bytes = uploaded_file.read()  # Read the image as bytes
            extracted_text = extract_text_from_image(image_bytes)
            st.write("Text extracted from the image:")
            st.write(extracted_text)
        elif file_extension == "txt":
            # Display TXT content
            txt_text = uploaded_file.read()
            text = txt_text
        elif file_extension == "mp3":
            # Recognize speech from the audio file
            recognized_text = recognize_speech(uploaded_file)
            st.subheader("Recognized Text from Audio:")
            st.write(recognized_text)
            text = recognized_text

        if text is not None:
            st.subheader("Text Extracted from Uploaded File:")
            st.write(text)

            # Count words in the text
            word_count = count_words(text)
            st.subheader(f"Word Count: {word_count} words")

            # Check if word count exceeds 1000
            if word_count > 1000:
                st.warning("Warning: The document contains more than 1000 words, which may be too large for translation.")
            else:
                # Add buttons for summarization and translation
                if st.button("Summarize Text"):
                    summarized_text = summarize_text(text)
                    st.subheader("Summarized Text:")
                    st.write(summarized_text)

                st.subheader('Select Language to Translate:')
                target_language = st.selectbox("Select target language:", list(language_mapping.values()))

                if st.button("Translate Text"):
                    # Translate the extracted text
                    try:
                        translated_text = translate_text(text, target_language_code)
                        st.subheader(f"Translated text ({target_language}):")
                        st.write(translated_text)
                    except Exception as e:
                        st.error(f"Translation error: {str(e)}")

                # Add button to convert translated text to speech and get the translated document
                if st.button("Convert to Speech and get Translated document"):
                    output_file = "translated_speech.mp3"
                    convert_text_to_speech(translated_text, output_file, language=target_language_code)

                    # Play the generated speech
                    audio_file = open(output_file, 'rb')
                    st.audio(audio_file.read(), format='audio/mp3')

                    # Play the generated speech (platform-dependent)
                    if os.name == 'posix':  # For Unix/Linux
                        os.system(f"xdg-open {output_file}")
                    elif os.name == 'nt':  # For Windows
                        os.system(f"start {output_file}")
                    else:
                        st.warning("Unsupported operating system")

                    # Provide a download link for the MP3 file
                    st.markdown(get_binary_file_downloader_html("Download Audio File", output_file, 'audio/mp3'), unsafe_allow_html=True)

                    # Convert the translated text to a Word document
                    word_output_file = "translated_text.docx"
                    convert_text_to_word_doc(translated_text, word_output_file)
                    # Provide a download link for the Word document
                    st.markdown(get_binary_file_downloader_html("Download Word Document", word_output_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'), unsafe_allow_html=True)

if __name__ == "__main__":
    main()
