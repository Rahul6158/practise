import streamlit as st
import os
import base64
import docx2txt
from googletrans import Translator as GoogleTranslator
from gtts import gTTS
import io
from docx import Document
from bs4 import BeautifulSoup
from PIL import Image
import pytesseract
import easyocr
import PyPDF2
from PIL import Image
from PyPDF2 import PdfFileWriter, PdfFileReader
import tempfile
from PyPDF2 import PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from fpdf import FPDF


language_mapping = {
    "en": "English",
    "hi": "Hindi",
    "gu": "Gujarati",
    "ta": "Tamil",
    "te": "Telugu",
    "kn": "Kannada",
    "mr": "Marathi",
    "bn": "Bengali",
}


# Function to extract text from a DOCX file
def process_docx_text(docx_file, skip_lists=True):
    # Extract text from the DOCX file
    if skip_lists:
        # Use custom function to remove lists
        text = process_docx_text_without_lists(docx_file)
    else:
        text = docx2txt.process(docx_file)
    return text

# Function to extract text from an uploaded image using Pytesseract
def extract_text_from_uploaded_image(image_bytes, language='eng'):
    try:
        image = Image.open(io.BytesIO(image_bytes))
        image = image.convert('RGB')
        text = pytesseract.image_to_string(image, lang=language)
        return text
    except Exception as e:
        return str(e)

# Custom function to remove lists from DOCX text
def process_docx_text_without_lists(docx_file):
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        if not paragraph.style.name.startswith('•'):
            text += paragraph.text + '\n'
    return text

# Function to extract text from a PDF file without lists
def process_pdf_text_without_lists(pdf_file):
    pdf_text = ""
    try:
        with st.spinner("Extracting text from PDF..."):
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
            for page_number in range(num_pages):
                page = pdf_reader.pages[page_number]
                pdf_text += page.extract_text()
    except Exception as e:
        st.error(f"Error processing PDF: {str(e)}")
    return pdf_text

# Function to extract text from a TXT file
def process_txt_file(txt_file):
    txt_text = txt_file.read()
    text = txt_text.decode('utf-8')
    return text

# Function to translate text using Google Translate
def translate_text_with_google(text, target_language):
    google_translator = GoogleTranslator()

    max_chunk_length = 500
    translated_text = ""

    for i in range(0, len(text), max_chunk_length):
        chunk = text[i:i + max_chunk_length]
        translated_chunk = google_translator.translate(chunk, dest=target_language).text
        translated_text += translated_chunk

    return translated_text

# Function to convert text to speech and save as an MP3 file
def convert_text_to_speech(text, output_file, language='en'):
    if text:
        supported_languages = list(language_mapping.keys())  # Add more supported languages as needed
        if language not in supported_languages:
            st.warning(f"Unsupported language: {language}")
            return

        tts = gTTS(text=text, lang=language)
        tts.save(output_file)

# Function to generate a download link for a file
def get_binary_file_downloader_html(link_text, file_path, file_format):
    with open(file_path, 'rb') as f:
        file_data = f.read()
    b64_file = base64.b64encode(file_data).decode()
    download_link = f'<a href="data:{file_format};base64,{b64_file}" download="{os.path.basename(file_path)}">{link_text}</a>'
    return download_link

# Function to convert translated text to a Word document
def convert_text_to_word_doc(text, output_file):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_file)

# Function to translate text with fallback to Google Translate on error
def translate_text_with_fallback(text, target_language):
    try:
        return translate_text_with_google(text, target_language)
    except Exception as e:
        st.warning(f"Google Translate error: {str(e)}")

# Function to convert translated text to a Word document
def convert_text_to_word_doc(text, output_file):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_file)

# Function to count words in the text
def count_words(text):
    words = text.split()
    return len(words)

# ... Import statements and function definitions ...

# Streamlit app main function
def main():
    st.image("jangirii.png", width=50)
    st.title("Text Translation and Conversion to Speech (English - other languages)")

    # Add a file uploader for DOCX, PDF, images
    uploaded_file = st.file_uploader("Upload a file", type=["docx", "pdf", "jpg", "jpeg", "png", "txt"])

    if uploaded_file is not None:
        file_extension = uploaded_file.name.split('.')[-1].lower()

        # Initialize text as None
        text = None

        if file_extension == "docx":
            # Display DOCX content
            text = process_docx_text_without_lists(uploaded_file)
        elif file_extension == "pdf":
            # Display PDF content
            pdf_text = process_pdf_text_without_lists(uploaded_file)
            text = pdf_text
        elif file_extension in ["jpg", "jpeg", "png"]:
            # Display image
            img = Image.open(uploaded_file)
            st.image(img, caption="Uploaded Image", use_column_width=True)

            # Extract text from the image using custom function
            image_bytes = uploaded_file.read()  # Read the image as bytes
            extracted_text = extract_text_from_uploaded_image(image_bytes)
            st.write("Text extracted from the image:")
            st.write(extracted_text)
        elif file_extension == "txt":
            # Display TXT content
            txt_text = uploaded_file.read()
            text = txt_text

        if text is not None:
            st.subheader("Text Extracted from Uploaded File:")
            st.write(text)

            # Count words in the text
            word_count = count_words(text)
            st.subheader(f"Word Count: {word_count} words")

            # Check if word count exceeds 5000
            if word_count > 15000:
                st.warning("Warning: The document contains more than 5000 words, which may be too large for translation.")
                return  # Exit the function if word count exceeds 5000

            st.subheader('Select Language to Translate : ')
            target_language = st.selectbox("Select target Regional language:", list(language_mapping.values()))

            # Check if text is not empty or None before attempting translation
            if text and len(text.strip()) > 0:
                # Add a button to trigger the translation and text-to-speech conversion
                if st.button("Translate - Convert to Speech and get Translated document"):
                    # Define target_language_code within this scope
                    target_language_code = [code for code, lang in language_mapping.items() if lang == target_language][0]

                    # Translate text using Google Translate
                    try:
                        translated_text = translate_text_with_google(text, target_language_code)
                    except Exception as e:
                        st.error(f"Google Translate error: {str(e)}")
                        translated_text = None

                    # Display translated text
                    if translated_text:
                        st.subheader(f"Translated text ({target_language}):")
                        st.write(translated_text)
                    else:
                        st.warning("Translation result is empty. Please check your input text.")

                    # Convert translated text to speech
                    output_file = "translated_speech.mp3"
                    convert_text_to_speech(translated_text, output_file, language=target_language_code)

                    # Play the generated speech
                    audio_file = open(output_file, 'rb')
                    st.audio(audio_file.read(), format='audio/mp3')

                    # Play the generated speech (platform-dependent)
                    if os.name == 'posix':  # For Unix/Linux
                        os.system(f"xdg-open {output_file}")
                    elif os.name == 'nt':  # For Windows
                        os.system(f"start {output_file}")
                    else:
                        st.warning("Unsupported operating system")

                    # Provide a download link for the MP3 file
                    st.markdown(get_binary_file_downloader_html("Download Audio File", output_file, 'audio/mp3'), unsafe_allow_html=True)

                    # Convert the translated text to a PDF document
                    pdf_output_file = "translated_text.pdf"
                    convert_text_to_pdf(translated_text, pdf_output_file)

                    # Provide a download link for the PDF document
                    st.markdown(get_binary_file_downloader_html("Download PDF Document", pdf_output_file, 'application/pdf'), unsafe_allow_html=True)

if __name__ == "__main__":
    main()
