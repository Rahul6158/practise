import streamlit as st
import os
import base64
import docx2txt
from translate import Translator
from gtts import gTTS
from docx import Document
from bs4 import BeautifulSoup
from PIL import Image
import fitz  # PyMuPDF

# Function to extract text from a DOCX file
def process_docx_text(docx_file, skip_lists=True):
    # Extract text from the DOCX file
    if skip_lists:
        # Use custom function to remove lists
        text = process_docx_text_without_lists(docx_file)
    else:
        text = docx2txt.process(docx_file)
    return text

# Custom function to remove lists from DOCX text
def process_docx_text_without_lists(docx_file):
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        if not paragraph.style.name.startswith('•'):
            text += paragraph.text + '\n'
    return text

# Function to extract text from a PDF file without lists
def process_pdf_text_without_lists(pdf_file):
    pdf_text = ""
    with st.spinner("Extracting text from PDF..."):
        pdf_bytes = pdf_file.read()
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_number in range(len(pdf_document)):
            page = pdf_document.load_page(page_number)
            pdf_text += page.get_text()
    return pdf_text

# Function to translate text using the translate library with a loop
def translate_text(text, target_language):
    translator = Translator(to_lang=target_language)
    max_chunk_length = 500
    translated_text = ""

    for i in range(0, len(text), max_chunk_length):
        chunk = text[i:i + max_chunk_length]
        translated_chunk = translator.translate(chunk)
        translated_text += translated_chunk

    return translated_text

# Function to translate text using the google translate library with a loop
def translate_text_google(text, target_language):
    translator = GoogleTranslator()
    max_chunk_length = 500
    translated_text = ""

    for i in range(0, len(text), max_chunk_length):
        chunk = text[i:i + max_chunk_length]
        translated_chunk = translator.translate(chunk, dest=target_language)
        translated_text += translated_chunk.text

    return translated_text

# Function to convert text to speech and save as an MP3 file
def convert_text_to_speech(text, output_file, language='en'):
    if text:
        tts = gTTS(text=text, lang=language)
        tts.save(output_file)

# Function to generate a download link for a file
def get_binary_file_downloader_html(link_text, file_path, file_format):
    with open(file_path, 'rb') as f:
        file_data = f.read()
    b64_file = base64.b64encode(file_data).decode()
    download_link = f'<a href="data:{file_format};base64,{b64_file}" download="{os.path.basename(file_path)}">{link_text}</a>'
    return download_link

# Function to convert translated text to a Word document
def convert_text_to_word_doc(text, output_file):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_file)

# Function to convert Word document to HTML
def convert_word_doc_to_html(docx_file):
    txt = docx2txt.process(docx_file)
    soup = BeautifulSoup(txt, 'html.parser')
    return soup.prettify()

def translate_text_with_fallback(text, target_language):
    try:
        return translate_text_mymemory(text, target_language)
    except Exception as e:
        st.warning(f"MyMemory translation error: {str(e)}")

    # If MyMemory fails, use Google Translate
    st.warning("Falling back to Google Translate...")
    try:
        return translate_text_google(text, target_language)
    except Exception as e:
        st.error(f"Google Translate error: {str(e)}")

    return "Translation failed. Please try again later."


# Function to count words in the text
def count_words(text):
    words = text.split()
    return len(words)

language_mapping = {
    "en": "English",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
    "it": "Italian",
    "pt": "Portuguese",
    "nl": "Dutch",
    "hi": "Hindi",
    "ja": "Japanese",
    "ko": "Korean",
    "zh-cn": "Simplified Chinese",
    "ru": "Russian",
    "ar": "Arabic",
    "th": "Thai",
    "tr": "Turkish",
    "pl": "Polish",
    "cs": "Czech",
    "sv": "Swedish",
    "da": "Danish",
    "fi": "Finnish",
    "el": "Greek",
    "hu": "Hungarian",
    "uk": "Ukrainian",
    "no": "Norwegian",
    "id": "Indonesian",
    "vi": "Vietnamese",
    "ro": "Romanian",
    "bn": "Bengali",
    "fa": "Persian",
    "iw": "Hebrew",
    "bg": "Bulgarian",
    "ca": "Catalan",
    "hr": "Croatian",
    "sr": "Serbian",
    "sk": "Slovak",
    "sl": "Slovenian",
    "lt": "Lithuanian",
    "lv": "Latvian",
    "et": "Estonian",
    "is": "Icelandic",
    "ga": "Irish",
    "sq": "Albanian",
    "mk": "Macedonian",
    "hy": "Armenian",
    "ka": "Georgian",
    "mt": "Maltese",
    "mr": "Marathi",
    "ta": "Tamil",
    "te": "Telugu",
    "ur": "Urdu",
    "ne": "Nepali",
    "si": "Sinhala",
    "km": "Khmer",
    "lo": "Lao",
    "my": "Burmese",
    "jw": "Javanese",
    "mn": "Mongolian",
    "zu": "Zulu",
    "xh": "Xhosa"
}

# Main Streamlit app
def main():
    st.image("jangirii.png", width=50)
    st.title("Text Translation and Conversion to Speech (English - other languages)")

    # Add a file uploader for DOCX, PDF, images
    uploaded_file = st.file_uploader("Upload a file", type=["docx", "pdf", "jpg", "jpeg", "png", "txt"])

    if uploaded_file is not None:
        file_extension = uploaded_file.name.split('.')[-1].lower()

        # Initialize text as None
        text = None
        output_file = "translated_speech.mp3"  # Define output_file here

        if file_extension == "docx":
            # Display DOCX content
            text = process_docx_text_without_lists(uploaded_file)
        elif file_extension == "pdf":
            # Display PDF content
            pdf_text = process_pdf_text_without_lists(uploaded_file)
            text = pdf_text
        elif file_extension in ["jpg", "jpeg", "png"]:
            # Display image
            img = Image.open(uploaded_file)
            st.image(img, caption="Uploaded Image", use_column_width=True)

            # Extract text from the image using custom function
            text = extract_text_from_image(img)
            st.write("Text extracted from the image:")
            st.write(text)
        elif file_extension == "txt":
            # Display TXT content
            txt_text = uploaded_file.read()
            text = txt_text

        if text is not None:
            st.subheader("Text Extracted from Uploaded File:")
            st.write(text)

            # Count words in the text
            word_count = count_words(text)
            st.subheader(f"Word Count: {word_count} words")

            # Check if word count exceeds 1000
            if word_count > 1000:
                st.warning("Warning: The document contains more than 1000 words, which may be too large for translation.")
                return  # Exit the function if word count exceeds 1000

            st.subheader('Select Language to Translate : ')
            target_language = st.selectbox("Select target language:", list(language_mapping.values()))

            # Check if the target language is in the mapping
            target_language_code = [code for code, lang in language_mapping.items() if lang == target_language][0]

            # Check if text is not empty or None before attempting translation
            if text and len(text.strip()) > 0:
                # Translate the extracted text
                try:
                    translated_text = translate_text_with_fallback(text, target_language_code)
                except Exception as e:
                    st.error(f"Translation error: {str(e)}")
                    translated_text = None
            else:
                st.warning("Input text is empty. Please check your document.")
                translated_text = None

            # Display translated text
            if translated_text:
                st.subheader(f"Translated text ({target_language}):")
                st.write(translated_text)
            else:
                st.warning("Translation result is empty. Please check your input text.")

            # Convert the translated text to speech
            if st.button("Convert to Speech and get Translated document"):
                # Use the output_file defined at the beginning
                convert_text_to_speech(translated_text, output_file, language=target_language_code)

                # Play the generated speech
                audio_file = open(output_file, 'rb')
                st.audio(audio_file.read(), format='audio/mp3')

                # Provide a download link for the MP3 file
                st.markdown(get_binary_file_downloader_html("Download Audio File", output_file, 'audio/mp3'), unsafe_allow_html=True)

                # Convert the translated text to a Word document
                word_output_file = "translated_text.docx"
                convert_text_to_word_doc(translated_text, word_output_file)
                # Provide a download link for the Word document
                st.markdown(get_binary_file_downloader_html("Download Word Document", word_output_file, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'), unsafe_allow_html=True)

if __name__ == "__main__":
    main()
